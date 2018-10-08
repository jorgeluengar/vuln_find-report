#########################################################################################################
#   Author: Jorge Luengo García
#   Purpose: Find SW vulnerabilities automatically from xls file
#   Description:
#   ./find_vulns.py <input_file.xls> <output_file_name> <VULNERS_API_KEY>
#   It will create a new xls file from input, including all cves found in vulners and nist database.
#   It will also create a word report containing all vulnerabilities and a description of each one.
#   Note: <output_file_name> will be used for XLS and DOCX files
#########################################################################################################

import xlrd
import xlwt
import sys
import vulners
import time
from docx import Document

if len(sys.argv) != 4:
    print("USAGE: ./find_vulns.py <input_file.xls> <output_file_name> <VULNERS_API_KEY>")
    exit(-1)

# Vulners Api KEY
vulners_api = vulners.Vulners(
    api_key=sys.argv[3])

# XLS input file
csvfilename = sys.argv[1]
outputName = sys.argv[2]

wbOutName = outputName + '.xls'
docxName = outputName + '.docx'

# File configuration
developer_column = 0
software_column = 1
version_column = 2
cve_vulners_column = 3
cve_cpe_column = 4
exploits_column = 5
date_column = 6

# Configure splitter
splitter = ", "

# To open XLS file
wb = xlrd.open_workbook(csvfilename)
sheet = wb.sheet_by_index(0)

# It creates a new Workbook to write results
wbOut = xlwt.Workbook(encoding='ascii')
wSheetOut = wbOut.add_sheet('My Worksheet')
wSheetOut.write(0, developer_column, label="Developer")
wSheetOut.write(0, software_column, label="Software")
wSheetOut.write(0, version_column, label="Version")
wSheetOut.write(0, cve_vulners_column, label="CVE vulners")
wSheetOut.write(0, cve_cpe_column, label="CVE CPE")
wSheetOut.write(0, exploits_column, label="Exploits")
wSheetOut.write(0, date_column, label="Research Date")

# Create Word Document
document = Document()
document.add_heading("Vulnerability Report " + time.strftime("%d/%m/%y"), 0)

# It takes all SW and its version from xls document
for i in range(1, sheet.nrows):

    software = sheet.cell_value(i, software_column)
    version = sheet.cell_value(i, version_column)
    developer = sheet.cell_value(i, developer_column)

    document.add_heading(
        'Software: ' + str(software) + " " + str(version), level=1
    )

    # It requests to vulners all vulnerabilities
    results = vulners_api.softwareVulnerabilities(software, str(version), 1000)
    results_cpe = vulners_api.cpeVulnerabilities("cpe:/a:" + str(developer) + ":" + str(software) + ":" + str(version))

    print("Working on:")
    print(str(i) + " -->  " + str(software) + " " + str(version))

    cve_vulners_list = ""
    document.add_heading('CVEs from Vulners', level=2)
    if str(results) == "{}":
        print("No vulnerability found.")
        document.add_paragraph("No vulnerability found in CPE.")
    else:
        cve_list = results["software"][0]["cvelist"]
        for cve in cve_list:
            print(cve)
            document.add_heading(str(cve), level=3)
            cve_vulners_list = cve_vulners_list + splitter + str(cve)

    cve_cpe_list = ""
    document.add_heading('CVEs from CPE', level=2)
    if str(results_cpe) == "{}":
        print("No vulnerability found in CPE.")
        document.add_paragraph("No vulnerability found in CPE.")
    else:
        cve_list = results_cpe["NVD"]
        print(cve_list)
        for cve in cve_list:
            print(cve["id"])
            print("SCORE: " + str(cve["cvss"]["score"]))

            # Write data to document
            document.add_heading(str(cve["id"]) + " --> " + "Score: " + str(cve["cvss"]["score"]), level=3)
            document.add_paragraph(str(cve["description"]))

            # Create CVE cpe list for XLS file
            cve_cpe_list = cve_cpe_list + splitter + str(cve["id"])

    wSheetOut.write(i, developer_column, label=developer)
    wSheetOut.write(i, software_column, label=software)
    wSheetOut.write(i, version_column, label=version)
    wSheetOut.write(i, cve_vulners_column, label=str(cve_vulners_list))
    wSheetOut.write(i, cve_cpe_column, label=str(cve_cpe_list))
    wSheetOut.write(i, date_column, label=time.strftime("%d/%m/%y"))

    print()

# Save XLS and DOCX documents
wbOut.save(wbOutName)
document.save(docxName)
